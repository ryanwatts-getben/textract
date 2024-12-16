# Standard library imports
import json
import logging
import os
import subprocess
import tempfile
import pickle
import re
import io
import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Optional, Tuple, Dict, Any
from pathlib import Path
import shutil
from pypdf import PdfReader
from docx.document import Document as DocxDocument
import csv
import torch

# Third-party imports
from flask import Flask, request, jsonify
from flask_cors import CORS
import boto3
from botocore.exceptions import ClientError
from dotenv import load_dotenv
import xml.etree.ElementTree as ET
from xml.etree.ElementTree import ParseError
from defusedxml import ElementTree as DefusedET
from flask_restx import Api, Resource

# LlamaIndex imports
from llama_index.core import (
    VectorStoreIndex,
    Document,
    load_index_from_storage,
)
from llama_index.llms.anthropic import Anthropic

# Local imports
from swagger_config import (
    api, scrape_ns, scrape_request, scrape_response, 
    scrape_multiple_response, error_response
)
from rag import create_index, preprocess_document
from disease_definition_generator import generate_multiple_definitions, get_embedding_model
from app_rag_disease_config import (
    ENV_PATH,
    AWS_UPLOAD_BUCKET_NAME,
    ALLOWED_ORIGINS,
    CORS_CONFIG,
    SUPPORTED_EXTENSIONS,
    DEFAULT_USER_ID,
    DEFAULT_PROJECT_ID,
    MAX_WORKERS,
    CUDA_MEMORY_FRACTION,
    CUDA_DEVICE,
    LLM_MODEL,
    QUERY_ENGINE_CONFIG,
    LOG_CONFIG,
    SERVER_CONFIG,
    DocumentConfig,
    ERROR_MESSAGES,
    TEMP_DIR_PREFIX,
    INDEX_CACHE_FILENAME,
    INDEX_METADATA_FILENAME
)
from scrape import scrape_medline_plus

# Load environment variables from .env file
load_dotenv(dotenv_path=ENV_PATH)

# Set up logging
logging.basicConfig(
    level=getattr(logging, LOG_CONFIG["level"]),
    format=LOG_CONFIG["format"]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Update CORS configuration
CORS(app, resources=CORS_CONFIG)

api.init_app(app)

def configure_cuda_memory_limit(fraction=CUDA_MEMORY_FRACTION, device=CUDA_DEVICE):
    if torch.cuda.is_available():
        torch.cuda.set_per_process_memory_fraction(fraction, device=device)
        torch.cuda.empty_cache()
        initial_memory = torch.cuda.memory_allocated(device=device)
        total_memory = torch.cuda.get_device_properties(device).total_memory
        logger.info(f"[rag] Initial GPU Memory Allocated: {initial_memory / 1e9:.2f} GB / {total_memory / 1e9:.2f} GB")
    else:
        logger.info("CUDA is not available.")

s3_client = boto3.client('s3')  # Initialize S3 client

def process_file(bucket_name, file_path, script_name):
    user_id, case_id = file_path.split('/')[:2]
    file_name = os.path.basename(file_path)
    file_name_without_extension = os.path.splitext(file_name)[0]
    message_body = {
        'file_name': file_name_without_extension,
        'file_path': file_path,
        'bucket': bucket_name,
        'case_id': case_id,
        'user_id': user_id,
    }
    logger.info(f"[app] Processing: {message_body}")
    # Construct the command to run the appropriate split script
    command = [
        "python", script_name,
        json.dumps(message_body)
    ]
    try:
        subprocess.run(command, check=True)
        logger.info(f"[app] Successfully processed {file_path}")
    except subprocess.CalledProcessError as e:
        logger.error(f"[app] Error processing {file_path}: {e}")
        raise

@app.route('/process', methods=['POST'])
def process():
    try:
        data = request.json
        bucket_name = data['bucket']
        records = data.get('records', [])
        bills = data.get('bills', [])
        accident_reports = data.get('reports', [])
        total_records = len(records)
        total_bills = len(bills)
        logger.info(f'[app] Total records: {total_records}, Total bills: {total_bills}')
        
        for accident_report in accident_reports:
            process_file(bucket_name, accident_report['file_path'], 'split_reports.py')
        for record in records:
            process_file(bucket_name, record['file_path'], 'split_records.py')
        for bill in bills:
            process_file(bucket_name, bill['file_path'], 'split_bills.py')
       
        logger.info(f"[app] Successfully processed {total_records} records and {total_bills} bills")
        return jsonify({"status": "success", "message": f"Successfully processed {total_records} records and {total_bills} bills"}), 200
    except Exception as e:
        logger.error(f"[app] Error occurred while processing records and bills: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500
    
def query_index(index: VectorStoreIndex, query_text: str) -> str:
    """Query the index with given text using Claude."""
    try:
        # Initialize Claude with the latest model
        llm = Anthropic(
            model=LLM_MODEL,
            api_key=os.getenv('ANTHROPIC_API_KEY')
        )

        # Create query engine with detailed configurations
        query_engine = index.as_query_engine(
            llm=llm,
            **QUERY_ENGINE_CONFIG
        )
        
        # Log the query for debugging
        logger.debug(f'[rag] Query Text: {query_text}')
        
        # Execute the query and get response
        response = query_engine.query(query_text)
        
        # Log the response for debugging
        logger.debug(f'[rag] Response Text: {response}')
        
        return response.response
        
    except Exception as e:
        logger.error(f"[app] Error querying index: {str(e)}")
        raise

def extract_text_from_file(content: bytes, file_extension: str) -> str:
    """
    Extract text from various file types.
    
    Args:
        content (bytes): The binary content of the file
        file_extension (str): The file extension (including the dot)
        
    Returns:
        str: The extracted text
    """
    try:
        if file_extension == '.txt':
            return content.decode('utf-8', errors='ignore')
            
        elif file_extension == '.pdf':
            with io.BytesIO(content) as pdf_buffer:
                reader = PdfReader(pdf_buffer)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text() or '')
                return '\n'.join(text)
                
        elif file_extension == '.docx':
            with io.BytesIO(content) as docx_buffer:
                doc = DocxDocument(docx_buffer)
                return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
                
        elif file_extension == '.csv':
            text_data = content.decode('utf-8', errors='ignore')
            csv_reader = csv.reader(text_data.splitlines())
            return '\n'.join(','.join(row) for row in csv_reader)
            
        elif file_extension in {'.xml', '.xsd'}:
            try:
                root = DefusedET.fromstring(content.decode('utf-8', errors='ignore'))
                text_content = []
                
                def extract_text_from_element(element, path=""):
                    current_path = f"{path}/{element.tag}" if path else element.tag
                    
                    # Add element text if present
                    if element.text and element.text.strip():
                        text_content.append(f"{current_path}: {element.text.strip()}")
                    
                    # Process attributes
                    for key, value in element.attrib.items():
                        text_content.append(f"{current_path}[@{key}]: {value}")
                    
                    # Process child elements
                    for child in element:
                        extract_text_from_element(child, current_path)
                
                extract_text_from_element(root)
                return '\n'.join(text_content)
                
            except Exception as xml_e:
                logger.error(f"[app] Error parsing XML/XSD content: {str(xml_e)}")
                return content.decode('utf-8', errors='ignore')
                
        elif file_extension == '.rrf':
            # RRF (Rich Release Format) files are typically pipe-delimited
            text_data = content.decode('utf-8', errors='ignore')
            return '\n'.join(line for line in text_data.splitlines())
            
        elif file_extension == '.json':
            try:
                json_data = json.loads(content.decode('utf-8', errors='ignore'))
                if isinstance(json_data, dict):
                    return '\n'.join(f"{k}: {v}" for k, v in json_data.items())
                elif isinstance(json_data, list):
                    return '\n'.join(str(item) for item in json_data)
                else:
                    return str(json_data)
            except json.JSONDecodeError:
                return content.decode('utf-8', errors='ignore')
        
        else:
            logger.warning(f"[app] Unsupported file extension: {file_extension}")
            return content.decode('utf-8', errors='ignore')
            
    except Exception as e:
        logger.error(f"[app] Error extracting text from {file_extension} file: {str(e)}")
        return ""

def process_single_document(s3_client, key: str) -> Optional[Document]:
    """Process a single document from S3."""
    try:
        logger.info(f"[app] Processing document: {key}")
        
        # Get the file from S3
        response = s3_client.get_object(Bucket=AWS_UPLOAD_BUCKET_NAME, Key=key)
        content = response['Body'].read()
        
        # Get file extension
        file_extension = os.path.splitext(key)[1].lower()
        logger.info(f"[app] File extension: {file_extension}")
        
        if file_extension not in SUPPORTED_EXTENSIONS:
            logger.warning(f"[app] Unsupported file extension: {file_extension}")
            return None
            
        # Extract text based on file type
        if file_extension == '.pdf':
            text = extract_text_from_pdf(content)
        elif file_extension == '.docx':
            text = extract_text_from_docx(content)
        elif file_extension == '.json':
            text = process_json_content(json.loads(content.decode('utf-8')))
        elif file_extension == '.xml' or file_extension == '.xsd':
            text = extract_text_from_xml(content)
        elif file_extension == '.csv':
            text = extract_text_from_csv(content)
        else:  # .txt and other text files
            text = content.decode('utf-8', errors='ignore')
            
        if not text:
            logger.warning(f"[app] No text content extracted from {key}")
            return None
            
        # Create document with metadata
        document = Document(
            text=text,
            metadata={
                'source': key,
                'file_type': file_extension,
                'timestamp': str(response.get('LastModified', '')),
                'size': response.get('ContentLength', 0)
            }
        )
        
        logger.info(f"[app] Successfully processed document: {key} (size: {len(text)} chars)")
        return document
        
    except Exception as e:
        logger.error(f"[app] Error processing document {key}: {str(e)}")
        return None

def create_or_update_index(user_id: str = DEFAULT_USER_ID, project_id: str = DEFAULT_PROJECT_ID, force_refresh: bool = False, max_workers: int = MAX_WORKERS) -> Tuple[VectorStoreIndex, str]:
    """
    Create or update an index for the specified user and project.
    """
    s3_client = boto3.client('s3')

    # Define the S3 paths
    input_prefix = f"{user_id}/{project_id}/input/"
    index_key = f"{user_id}/{project_id}/index.pkl"

    logger.info(f"[app] Looking for documents in s3://{AWS_UPLOAD_BUCKET_NAME}/{input_prefix}")

    temp_dir = tempfile.mkdtemp(prefix=TEMP_DIR_PREFIX)
    index_cache_path = os.path.join(temp_dir, INDEX_CACHE_FILENAME)

    try:
        # Check if index exists in S3
        try:
            s3_client.head_object(Bucket=AWS_UPLOAD_BUCKET_NAME, Key=index_key)
            index_exists = True
            logger.info(f"[app] Index found at s3://{AWS_UPLOAD_BUCKET_NAME}/{index_key}")
        except ClientError as e:
            # Index not found
            index_exists = False
            logger.info(f"[app] No existing index found at s3://{AWS_UPLOAD_BUCKET_NAME}/{index_key}")

        if index_exists and not force_refresh:
            # Load existing index from S3
            logger.info(f"[app] Loading existing index from s3://{AWS_UPLOAD_BUCKET_NAME}/{index_key}")
            with open(index_cache_path, 'wb') as f:
                s3_client.download_fileobj(AWS_UPLOAD_BUCKET_NAME, index_key, f)
            with open(index_cache_path, 'rb') as f:
                index = pickle.load(f)
            return index, 'Existing index loaded successfully'

        # Create new index or refresh existing one
        logger.info('[app] Creating new index')
        documents = []

        # List all objects in the input prefix
        try:
            paginator = s3_client.get_paginator('list_objects_v2')
            pages = paginator.paginate(Bucket=AWS_UPLOAD_BUCKET_NAME, Prefix=input_prefix)
            
            # Collect all valid files
            files_to_process = set()  # Using set to avoid duplicates
            all_files = []
            
            for page in pages:
                if 'Contents' not in page:
                    logger.warning(f"[app] No contents found in page for prefix: {input_prefix}")
                    continue
                    
                for obj in page.get('Contents', []):
                    key = obj.get('Key')
                    if not key:
                        continue
                        
                    all_files.append(key)
                    file_extension = os.path.splitext(key)[1].lower()
                    
                    # Skip duplicate files (e.g., icd10cm vs icd-10-cm)
                    base_name = os.path.basename(key).lower()
                    if any(x in base_name for x in ['icd10cm-', 'icd10-']):
                        if any(x in files_to_process for x in [key.replace('icd10cm-', 'icd-10-cm-'), 
                                                             key.replace('icd10-', 'icd-10-cm-')]):
                            logger.info(f"[app] Skipping duplicate file: {key}")
                            continue
                    
                    if file_extension in SUPPORTED_EXTENSIONS:
                        files_to_process.add(key)
                        logger.info(f"[app] Found valid file to process: {key}")
                    else:
                        logger.warning(f"[app] Skipping unsupported file: {key} (extension: {file_extension})")
            
            logger.info(f"[app] Found {len(all_files)} total files, {len(files_to_process)} valid files to process")
            logger.debug(f"[app] All files found: {all_files}")
            
            if not files_to_process:
                logger.error(f"[app] No valid files found in prefix: {input_prefix}")
                logger.error(f"[app] Supported extensions are: {SUPPORTED_EXTENSIONS}")
                raise ValueError(ERROR_MESSAGES["no_valid_files"])

        except ClientError as e:
            logger.error(f"[app] Error listing objects in bucket: {str(e)}")
            raise ValueError(f"Error accessing S3: {str(e)}")

        # Process files in parallel
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_key = {
                executor.submit(process_single_document, s3_client, key): key
                for key in files_to_process
            }
            for future in as_completed(future_to_key):
                key = future_to_key[future]
                try:
                    doc = future.result()
                    if doc:
                        documents.append(doc)
                        logger.info(f'[app] Added document from {os.path.basename(key)}')
                    else:
                        logger.warning(f'[app] No content extracted from {key}')
                except Exception as e:
                    logger.error(f'[app] Error processing file {key}: {e}')

        if not documents:
            logger.error('[app] No documents were successfully processed')
            raise ValueError(ERROR_MESSAGES["no_processed_docs"])

        # Create index
        logger.info(f'[app] Creating new index with {len(documents)} documents')
        embed_model, _ = get_embedding_model()
        index = create_index(
            documents=documents,
            embed_model=embed_model,
            batch_size=DocumentConfig.INDEX_BATCH_SIZE,
            chunk_size=DocumentConfig.CHUNK_SIZE
        )

        # Cache the index locally
        logger.info(f"[app] Saving index to local cache: {index_cache_path}")
        with open(index_cache_path, 'wb') as f:
            pickle.dump(index, f)

        # Upload the index to S3
        logger.info(f"[app] Uploading index to S3: s3://{AWS_UPLOAD_BUCKET_NAME}/{index_key}")
        s3_client.upload_file(index_cache_path, AWS_UPLOAD_BUCKET_NAME, index_key)
        logger.info(f"[app] Index upload complete. Local copy at: {index_cache_path}, S3 copy at: s3://{AWS_UPLOAD_BUCKET_NAME}/{index_key}")

        return index, 'New index created and uploaded successfully'

    finally:
        # Clean up temporary directory
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            logger.info(f"[app] Cleaned up temporary directory: {temp_dir}")

def preprocess_document(document_text: str) -> str:
    """Preprocess a single document's text."""
    return document_text.lower()

class ProgressPercentage:
    def __init__(self, filename):
        self._filename = filename
        self._size = float(os.path.getsize(filename))
        self._seen_so_far = 0
        self._lock = threading.Lock()
        self.progress = 0

    def __call__(self, bytes_amount):
        with self._lock:
            self._seen_so_far += bytes_amount
            self.progress = (self._seen_so_far / self._size) * 100
            logger.info(f'[app] Upload progress: {self.progress:.2f}%')

@app.route('/index', methods=['POST'])
def index_route():
    """Endpoint to create or update the index."""
    try:
        # Extract parameters from the request
        data = request.get_json()
        user_id = data.get('userId', '00000')
        project_id = data.get('projectId', '22222')
        force_refresh = str(data.get('force_refresh', 'false')).lower() == 'true'
        
        logger.info(f"[app] Received index request - user_id: {user_id}, project_id: {project_id}, force_refresh: {force_refresh}")
        
        # Configure CUDA memory limit
        configure_cuda_memory_limit(0.5, device=0)
        
        # Call the index creation function
        index, status = create_or_update_index(
            user_id=user_id,
            project_id=project_id,
            force_refresh=force_refresh
        )
        
        # Return a success response
        return jsonify({
            'status': 'success',
            'message': status
        }), 200
    except Exception as e:
        logger.error(f"[app] Error in /index route: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/query', methods=['POST'])
def query():
    try:
        data = request.json
        force_refresh = data.get('force_refresh', False)
        
        # Handle both naming conventions
        query_text = data.get('query_text') or data.get('query', '')
        project_id = data.get('project_id') or data.get('projectId', '')
        user_id = data.get('user_id') or data.get('userId', '')
        
        if not all([query_text, project_id, user_id]):
            missing = []
            if not query_text: missing.append('query text')
            if not project_id: missing.append('project ID')
            if not user_id: missing.append('user ID')
            return jsonify({
                'status': 'error',
                'message': ERROR_MESSAGES["missing_query_fields"].format(fields=", ".join(missing))
            }), 400

        # Get or create index
        try:
            index, _ = create_or_update_index(user_id, project_id, force_refresh)
            
            # Query the index
            response_text = query_index(index, query_text)
            if not response_text or response_text == "Empty Response":
                return jsonify({
                    'status': 'error',
                    'message': 'No relevant information found for your query.'
                }), 404

            return jsonify({'status': 'success', 'response': response_text}), 200
            
        except ValueError as ve:
            return jsonify({
                'status': 'error',
                'message': str(ve)
            }), 404
            
    except Exception as e:
        logger.exception(f'[app] Error processing query: {str(e)}')
        return jsonify({'status': 'error', 'message': str(e)}), 500

def process_pdf_page(page, page_num, total_pages):
    """Process a single PDF page."""
    try:
        page_text = page.extract_text()
        if page_text:
            text = page_text.strip()
            logger.info(f'[app] Successfully extracted page {page_num}/{total_pages} with {len(text)} characters')
            return text
        else:
            logger.warning(f'[app] Empty text extracted from page {page_num}/{total_pages}')
            return ""
    except Exception as e:
        logger.error(f'[app] Error extracting text from page {page_num}: {str(e)}')
        return ""

def read_pdf_with_fallbacks(file_path):
    """
    Attempt to read PDF text using multiple methods with fallbacks.
    Returns the best result or None if all methods fail.
    """
    text_results = []
    
    # Method 1: PyPDF
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            num_pages = len(pdf_reader.pages)
            logger.info(f'[app] PDF has {num_pages} pages')
            
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text.strip() + '\n'
                        logger.info(f'[app] Successfully extracted page {i+1}/{num_pages} with {len(page_text)} characters')
                except Exception as e:
                    logger.warning(f'[app] PyPDF failed on page {i+1}: {str(e)}')
                    
        if text.strip():
            text_results.append(('pypdf', text))
    except Exception as e:
        logger.error(f'[app] PyPDF extraction failed: {str(e)}')

    # # Method 2: pdfplumber
    # try:
    #     text = ""
    #     with pdfplumber.open(file_path) as pdf:
    #         for i, page in enumerate(pdf.pages):
    #             try:
    #                 page_text = page.extract_text()
    #                 if page_text:
    #                     text += page_text.strip() + '\n'
    #                     logger.info(f'[app] pdfplumber: Extracted page {i+1} with {len(page_text)} characters')
    #             except Exception as e:
    #                 logger.warning(f'[app] pdfplumber failed on page {i+1}: {str(e)}')
                    
    #     if text.strip():
    #         text_results.append(('pdfplumber', text))
    # except Exception as e:
    #     logger.error(f'[app] pdfplumber extraction failed: {str(e)}')

    # # Method 3: pdfminer
    # try:
    #     text = pdfminer_extract_text(file_path)
    #     if text.strip():
    #         text_results.append(('pdfminer', text))
    #         logger.info(f'[app] pdfminer: Extracted {len(text)} characters')
    # except Exception as e:
    #     logger.error(f'[app] pdfminer extraction failed: {str(e)}')

    # # Choose the best result
    # if text_results:
    #     # Sort by text length (assuming longer text is better)
    #     text_results.sort(key=lambda x: len(x[1]), reverse=True)
    #     method, text = text_results[0]
    #     logger.info(f'[app] Using {method} result with {len(text)} characters')
        return clean_text(text)
    
    return None

def clean_text(text):
    """Clean and normalize extracted text."""
    if not text:
        return None
        
    # Remove null bytes and control characters
    text = ''.join(char if ord(char) >= 32 or char in '\n\t' else ' ' for char in text)
    
    # Normalize whitespace
    text = ' '.join(text.split())
    
    # Remove any remaining problematic characters
    text = text.replace('\x00', '')
    
    # Remove very short lines (likely garbage)
    lines = [line for line in text.split('\n') if len(line.strip()) > DocumentConfig.PDF_MIN_LINE_LENGTH]
    text = '\n'.join(lines)
    
    return text if text.strip() else None

def read_file_content(file_path):
    """Reads and validates file content."""
    try:
        file_lower = file_path.lower()
        
        if file_lower.endswith('.pdf'):
            logger.info(f'[app] Reading PDF file: {file_path}')
            text = read_pdf_with_fallbacks(file_path)
            
        elif file_lower.endswith('.json'):
            logger.info(f'[app] Reading JSON file: {file_path}')
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                text = process_json_content(data)
                
        elif file_lower.endswith('.xml'):
            logger.info(f'[app] Reading XML file: {file_path}')
            with open(file_path, 'r', encoding='utf-8') as file:
                xml_content = file.read()
            xsd_content = None
            # Check for corresponding XSD file
            xsd_file_path = os.path.splitext(file_path)[0] + '.xsd'
            if os.path.exists(xsd_file_path):
                logger.info(f'[app] Found XSD schema: {xsd_file_path}')
                with open(xsd_file_path, 'r', encoding='utf-8') as xsd_file:
                    xsd_content = xsd_file.read()
            text = parse_xml_content(xml_content, xsd_content)
            
        elif file_lower.endswith('.xsd'):
            logger.info(f'[app] Reading XSD file: {file_path} (treating as XML)')
            with open(file_path, 'r', encoding='utf-8') as file:
                xml_content = file.read()
            text = parse_xml_content(xml_content)
            
        else:  # Text files
            logger.info(f'[app] Reading text file: {file_path}')
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

        # Validate and clean the text
        text = clean_text(text)
        
        if not text:
            logger.warning(f'[app] No valid content extracted from {file_path}')
            return None
            
        logger.info(f'[app] Successfully extracted {len(text)} characters from {file_path}')
        return text
        
    except Exception as e:
        logger.error(f"[app] Error reading file {file_path}: {e}")
        return None

def sanitize_filename(filename):
    """
    Sanitizes the filename to remove or replace characters not allowed in file systems.
    """
    # Replace invalid characters with an underscore
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

# Ensure that the index is initialized or loaded
# You might need to adjust this depending on your application flow
index = None  # Placeholder for the index variable

def parse_xml_content(xml_content: str, xsd_content: Optional[str] = None) -> str:
    """
    Parse XML content and optionally validate against XSD schema.
    """
    try:
        # Parse XML securely using DefusedET
        root = DefusedET.fromstring(xml_content)
        
        # If XSD content is provided, validate XML
        if xsd_content:
            logger.info("[app] Validating XML against XSD schema")
            # Parse XSD content securely
            xsd_root = DefusedET.fromstring(xsd_content)
            xml_schema = ET.XMLSchema(xsd_root)
            if not xml_schema.validate(root):
                logger.warning(f"[app] XML validation failed against XSD schema")
                # Optionally, raise an error or log details
                # xml_schema.assertValid(root)
        
        # Rest of the function remains the same
        text_content = []

        def process_element(element, path=""):
            """Recursively process XML elements."""
            current_path = f"{path}/{element.tag}" if path else element.tag

            # Add element text if present
            if element.text and element.text.strip():
                text_content.append(f"{current_path}: {element.text.strip()}")

            # Process attributes
            for key, value in element.attrib.items():
                text_content.append(f"{current_path}[@{key}]: {value}")

            # Process child elements
            for child in element:
                process_element(child, current_path)

        process_element(root)
        return '\n'.join(text_content)

    except ET.ParseError as e:
        logger.error(f"[app] XML parsing error: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"[app] Error processing XML content: {str(e)}")
        raise

def process_json_content(data):
    """Process JSON data and convert it to text."""
    if isinstance(data, dict):
        return '\n'.join(f"{k}: {v}" for k, v in data.items() if v)
    elif isinstance(data, list):
        return '\n'.join(str(item) for item in data if item)
    else:
        return str(data)

@app.route('/generate-multiple-disease-definitions', methods=['POST'])
def generate_definitions():
    """
    Endpoint to generate multiple disease definitions.
    Expects JSON payload: {"diseaseNames": ["disease1", "disease2", ...]}
    """
    try:
        logger.info("[app] Received request to generate disease definitions")
        # Validate request
        request_data = request.get_json()
        logger.info(f"[app] Request data received: {request_data}")
        
        if not request_data or 'diseaseNames' not in request_data:
            logger.warning("[app] Invalid request: missing diseaseNames")
            return jsonify({
                'error': ERROR_MESSAGES["missing_disease_names"]
            }), 400

        disease_names = request_data['diseaseNames']
        
        # Generate definitions using the imported function
        results = generate_multiple_definitions(disease_names)
        
        logger.info("[app] Definitions generated successfully")
        return jsonify(results)

    except ValueError as ve:
        logger.error(f"[app] Validation error: {str(ve)}")
        return jsonify({'error': str(ve)}), 400
    except Exception as e:
        logger.error(f"[app] Server error: {str(e)}")
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.errorhandler(404)
def not_found(error):
    logger.warning(f"[app] 404 Not Found: {request.path}")
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def server_error(error):
    logger.error(f"[app] 500 Internal Server Error: {error}")
    return jsonify({'error': 'Internal server error'}), 500

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content."""
    try:
        with io.BytesIO(content) as pdf_file:
            reader = PdfReader(pdf_file)
            text = []
            for page in reader.pages:
                text.append(page.extract_text() or '')
            return '\n'.join(text)
    except Exception as e:
        logger.error(f"[app] Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content."""
    try:
        with io.BytesIO(content) as docx_file:
            doc = DocxDocument(docx_file)
            return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        logger.error(f"[app] Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_xml(content: bytes) -> str:
    """Extract text from XML content."""
    try:
        root = DefusedET.fromstring(content.decode('utf-8', errors='ignore'))
        text_content = []
        
        def extract_text_from_element(element, path=""):
            current_path = f"{path}/{element.tag}" if path else element.tag
            
            # Add element text if present
            if element.text and element.text.strip():
                text_content.append(f"{current_path}: {element.text.strip()}")
            
            # Process attributes
            for key, value in element.attrib.items():
                text_content.append(f"{current_path}[@{key}]: {value}")
            
            # Process child elements
            for child in element:
                extract_text_from_element(child, current_path)
        
        extract_text_from_element(root)
        return '\n'.join(text_content)
        
    except Exception as e:
        logger.error(f"[app] Error extracting text from XML: {str(e)}")
        return ""

def extract_text_from_csv(content: bytes) -> str:
    """Extract text from CSV content."""
    try:
        text_data = content.decode('utf-8', errors='ignore')
        csv_reader = csv.reader(text_data.splitlines())
        return '\n'.join(','.join(str(cell) for cell in row) for row in csv_reader)
    except Exception as e:
        logger.error(f"[app] Error extracting text from CSV: {str(e)}")
        return ""

def process_json_content(data: Any) -> str:
    """Process JSON data and convert it to text."""
    try:
        if isinstance(data, dict):
            return '\n'.join(f"{k}: {v}" for k, v in data.items() if v is not None)
        elif isinstance(data, list):
            return '\n'.join(str(item) for item in data if item is not None)
        else:
            return str(data)
    except Exception as e:
        logger.error(f"[app] Error processing JSON content: {str(e)}")
        return ""

@scrape_ns.route('/')
class ScrapeDiseaseInfo(Resource):
    @scrape_ns.expect(scrape_request)
    @scrape_ns.response(200, 'Success', scrape_multiple_response)
    @scrape_ns.response(400, 'Bad Request', error_response)
    @scrape_ns.response(500, 'Server Error', error_response)
    def post(self):
        """
        Scrape disease information from MedlinePlus.
        
        Accepts two formats:
        1. Array format:
        {
            "diseaseUrl": [
                ["http://www.nlm.nih.gov/medlineplus/abdominalpain.html", "Abdominal Pain"]
            ]
        }
        
        2. Single URL format:
        {
            "diseaseUrl": "http://www.nlm.nih.gov/medlineplus/abdominalpain.html"
        }
        """
        try:
            # Log the raw request data for debugging
            raw_data = request.get_data()
            logger.info(f"[app] Raw request data: {raw_data}")
            
            # Get JSON data and log it
            data = request.get_json(force=True)
            logger.info(f"[app] Parsed request data: {data}")
            
            if not data:
                logger.error("[app] No JSON data in request body")
                return {
                    'error': 'No JSON data in request body'
                }, 400
                
            if 'diseaseUrl' not in data:
                logger.error("[app] Missing diseaseUrl in request body")
                return {
                    'error': 'Missing diseaseUrl in request body'
                }, 400
                
            disease_url_data = data['diseaseUrl']
            
            # Handle single URL format
            if isinstance(disease_url_data, str):
                logger.info(f"[app] Processing single URL: {disease_url_data}")
                result = scrape_medline_plus(disease_url_data)
                if result is not None:
                    return result
                else:
                    return {
                        'error': 'Failed to scrape content from the provided URL'
                    }, 500
            
            # Handle array format
            elif isinstance(disease_url_data, list):
                results = {}
                for url_pair in disease_url_data:
                    if not isinstance(url_pair, list) or len(url_pair) != 2:
                        logger.error(f"[app] Invalid URL pair format: {url_pair}")
                        return {
                            'error': 'Each item in diseaseUrl array must be a pair of [URL, name]'
                        }, 400
                        
                    url, disease_name = url_pair
                    logger.info(f"[app] Scraping data for {disease_name} from {url}")
                    
                    result = scrape_medline_plus(url)
                    
                    if result is not None:
                        results[disease_name] = result
                    else:
                        logger.error(f"[app] Failed to scrape content for {disease_name} from {url}")
                        results[disease_name] = {'error': f'Failed to scrape content for {disease_name}'}
                
                if not results:
                    return {
                        'error': 'Failed to scrape content from any of the provided URLs'
                    }, 500
                    
                return results
            
            else:
                logger.error("[app] Invalid diseaseUrl format")
                return {
                    'error': 'diseaseUrl must be either a string URL or an array of [URL, name] pairs'
                }, 400
            
        except Exception as e:
            logger.error(f"[app] Error processing scrape request: {str(e)}")
            return {
                'error': f'Internal server error: {str(e)}'
            }, 500

if __name__ == '__main__':
    logger.info(f"[app] Starting Flask app on {SERVER_CONFIG['host']}:{SERVER_CONFIG['port']}")
    app.run(host=SERVER_CONFIG['host'], port=SERVER_CONFIG['port'])